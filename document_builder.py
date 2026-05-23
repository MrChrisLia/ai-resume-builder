"""
document_builder.py — Resume & cover letter document generation.

Templates: modern | classic | minimal | japanese | taiwanese
Formats:   .docx | .pdf | .md
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import weasyprint
import os


# ── Shared helpers ─────────────────────────────────────────────────────────────

def _lang_str(l):
    """Format a language dict as a display string, including certificate if present."""
    name  = l.get('language', '')
    prof  = l.get('proficiency', '')
    cert  = l.get('certificate', '').strip()
    if prof and cert:
        return f"{name} ({prof}, {cert})"
    elif prof:
        return f"{name} ({prof})"
    elif cert:
        return f"{name} — {cert}"
    return name


def _set_font(run, size_pt, bold=False, italic=False, color=None, font_name=None):
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name


def _add_bottom_border(para, color_hex='1A1A2E'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_top_border(para, color_hex='CCCCCC'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '4')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), color_hex)
    pBdr.append(top)
    pPr.append(pBdr)


def _set_margins(doc, top=0.5, bottom=0.5, left=0.75, right=0.75):
    for section in doc.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)


def _contact_str(data):
    parts = [data.get('email',''), data.get('phone',''), data.get('location',''),
             data.get('linkedin',''), data.get('github','')]
    return '  |  '.join(p for p in parts if p)


def _photo_stream(data):
    """Return a BytesIO of the profile photo, or None."""
    photo = data.get('photo', '')
    if not photo:
        return None
    try:
        import base64, io as _io
        raw = photo.split(',', 1)[-1] if ',' in photo else photo
        return _io.BytesIO(base64.b64decode(raw + '=='))
    except Exception:
        return None


def _add_photo_header(doc, data, name_size, name_color,
                      contact_size, contact_color, font_name=None, center=True):
    """Add a name+contact header, placing photo in a side cell if present."""
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE

    name    = data.get('name', '')
    contact = _contact_str(data)
    stream  = _photo_stream(data)

    if stream:
        table = doc.add_table(rows=1, cols=2)
        # Remove all cell borders
        for cell in table._cells:
            tc    = cell._tc
            tcPr  = tc.get_or_add_tcPr()
            bdr   = _OE('w:tcBorders')
            for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                tag = _OE(f'w:{edge}')
                tag.set(_qn('w:val'), 'none')
                bdr.append(tag)
            tcPr.append(bdr)

        left  = table.cell(0, 0)
        right = table.cell(0, 1)

        # Fix column widths: text ~5 in, photo ~1.3 in
        for i, width in enumerate([Inches(5.0), Inches(1.4)]):
            table.columns[i].width = width

        # Name in left cell
        pn = left.paragraphs[0]
        if center: pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(pn.add_run(name), name_size, bold=True, color=name_color, font_name=font_name)
        if contact:
            pc = left.add_paragraph()
            if center: pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pc.paragraph_format.space_before = Pt(2)
            _set_font(pc.add_run(contact), contact_size, color=contact_color, font_name=font_name)

        # Photo in right cell
        pr = right.paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pr.paragraph_format.space_before = Pt(2)
        run = pr.add_run()
        try:
            run.add_picture(stream, height=Inches(1.1))
        except Exception:
            pass
    else:
        p = doc.add_paragraph()
        if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(name), name_size, bold=True, color=name_color, font_name=font_name)
        if contact:
            p2 = doc.add_paragraph()
            if center: p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_before = Pt(2)
            _set_font(p2.add_run(contact), contact_size, color=contact_color, font_name=font_name)


# ── Language helpers ───────────────────────────────────────────────────────────

def _present_label(language: str) -> str:
    if language == 'japanese':  return '現在に至る'
    if language == 'taiwanese': return '至今'
    return 'Present'


def _section_labels(language: str, template: str = 'modern') -> dict:
    """Return section heading strings for the given content language."""
    if language == 'japanese':
        # JP_SECTIONS / TW_SECTIONS are defined later in this module but resolved at call-time
        return {
            'summary': '志望動機', 'experience': '職歴', 'education': '学歴',
            'skills': '特技・スキル', 'languages': '語学力',
            'projects': '主なプロジェクト', 'certifications': '免許・資格',
        }
    if language == 'taiwanese':
        return {
            'summary': '個人簡介', 'experience': '工作經歷', 'education': '學歷',
            'skills': '專業技能', 'languages': '語言能力',
            'projects': '專案經歷', 'certifications': '證照與認證',
        }
    # English — label set depends on which template format is being used
    if template == 'japanese':
        return {
            'summary': 'Motivation / Summary', 'experience': 'Work History',
            'education': 'Education', 'skills': 'Special Skills',
            'languages': 'Languages', 'projects': 'Projects',
            'certifications': 'Licenses & Certifications',
        }
    return {
        'summary': 'Professional Summary', 'experience': 'Experience',
        'education': 'Education', 'skills': 'Skills', 'languages': 'Languages',
        'projects': 'Projects', 'certifications': 'Certifications / Professional Licenses',
    }


# ══════════════════════════════════════════════════════════════════════════════
#  MODERN TEMPLATE
# ══════════════════════════════════════════════════════════════════════════════

DARK = RGBColor(0x1A, 0x1A, 0x2E)
MID  = RGBColor(0x44, 0x44, 0x44)
LITE = RGBColor(0x77, 0x77, 0x77)
BLUE = RGBColor(0x22, 0x44, 0x88)


def _modern_section_heading(doc, title):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(3)
    run = para.add_run(title.upper())
    _set_font(run, 10, bold=True, color=DARK)
    _add_bottom_border(para, '1A1A2E')
    return para


def _modern_entry_header(doc, left_text, right_text=''):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(5)
    para.paragraph_format.space_after  = Pt(1)
    _set_font(para.add_run(left_text), 10, bold=True, color=DARK)
    if right_text:
        para.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        _set_font(para.add_run(f'\t{right_text}'), 9, color=LITE)
    return para


def _build_modern_docx(data: dict, output_path: str, language: str = 'english'):
    doc = Document()
    _set_margins(doc)
    doc.styles['Normal'].paragraph_format.space_after  = Pt(0)
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    lbl = _section_labels(language, 'modern')

    _add_photo_header(doc, data, 22, DARK, 9, MID)

    if data.get('summary'):
        _modern_section_heading(doc, lbl['summary'])
        p = doc.add_paragraph(data['summary'])
        p.runs[0].font.size = Pt(10)

    if data.get('experience'):
        _modern_section_heading(doc, lbl['experience'])
        for exp in data['experience']:
            end = exp.get('end_date') or _present_label(language)
            _modern_entry_header(doc, exp.get('company',''), f"{exp.get('start_date','')} – {end}")
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _set_font(p.add_run(exp.get('title','')), 10, italic=True, color=MID)
            for bullet in exp.get('bullets', []):
                bp = doc.add_paragraph(style='List Bullet')
                bp.paragraph_format.left_indent = Inches(0.2)
                bp.paragraph_format.space_after  = Pt(1)
                _set_font(bp.add_run(bullet), 10)

    if data.get('education'):
        _modern_section_heading(doc, lbl['education'])
        for edu in data['education']:
            _modern_entry_header(doc, edu.get('school',''), edu.get('graduation',''))
            deg = ', '.join(p for p in [edu.get('degree',''), edu.get('field','')] if p)
            if edu.get('gpa'): deg += f'  |  GPA: {edu["gpa"]}'
            if deg:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                _set_font(p.add_run(deg), 10, italic=True, color=MID)

    if data.get('skills'):
        _modern_section_heading(doc, lbl['skills'])
        for cat in data['skills']:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _set_font(p.add_run(f"{cat.get('category','')}: "), 10, bold=True)
            _set_font(p.add_run(', '.join(cat.get('items',[]))), 10)

    langs = [l for l in data.get('languages', []) if l.get('language')]
    if langs:
        _modern_section_heading(doc, lbl['languages'])
        p = doc.add_paragraph()
        _set_font(p.add_run(', '.join(_lang_str(l) for l in langs)), 10)

    if data.get('projects'):
        _modern_section_heading(doc, lbl['projects'])
        for proj in data['projects']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            _set_font(p.add_run(proj.get('name','')), 10, bold=True, color=DARK)
            if proj.get('technologies'):
                _set_font(p.add_run(f"  |  {proj['technologies']}"), 10, color=BLUE)
            if proj.get('description'):
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(2)
                _set_font(p2.add_run(proj['description']), 10)

    certs = [c for c in data.get('certifications', []) if c]
    if certs:
        _modern_section_heading(doc, lbl['certifications'])
        for cert in certs:
            bp = doc.add_paragraph(style='List Bullet')
            bp.paragraph_format.left_indent = Inches(0.2)
            bp.paragraph_format.space_after  = Pt(1)
            _set_font(bp.add_run(cert), 10)

    doc.save(output_path)


# ══════════════════════════════════════════════════════════════════════════════
#  CLASSIC TEMPLATE
# ══════════════════════════════════════════════════════════════════════════════

BLK = RGBColor(0x00, 0x00, 0x00)
GRY = RGBColor(0x33, 0x33, 0x33)


def _classic_section_heading(doc, title):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(3)
    run = para.add_run(title.upper())
    _set_font(run, 11, bold=True, color=BLK, font_name='Cambria')
    run.underline = True
    return para


def _build_classic_docx(data: dict, output_path: str, language: str = 'english'):
    doc = Document()
    _set_margins(doc, top=0.75, bottom=0.75)
    doc.styles['Normal'].paragraph_format.space_after  = Pt(0)
    lbl = _section_labels(language, 'classic')

    _add_photo_header(doc, data, 18, BLK, 10, GRY, font_name='Cambria')

    if data.get('summary'):
        _classic_section_heading(doc, lbl['summary'])
        p = doc.add_paragraph(data['summary'])
        _set_font(p.runs[0], 10, font_name='Cambria')

    if data.get('experience'):
        _classic_section_heading(doc, lbl['experience'])
        for exp in data['experience']:
            end = exp.get('end_date') or _present_label(language)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
            _set_font(p.add_run(f"{exp.get('company','')}"), 10, bold=True, font_name='Cambria')
            _set_font(p.add_run(f"\t{exp.get('start_date','')} – {end}"), 10, font_name='Cambria')
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(2)
            _set_font(p2.add_run(exp.get('title','')), 10, italic=True, font_name='Cambria')
            for bullet in exp.get('bullets', []):
                bp = doc.add_paragraph()
                bp.paragraph_format.left_indent = Inches(0.2)
                bp.paragraph_format.space_after  = Pt(1)
                _set_font(bp.add_run(f'— {bullet}'), 10, font_name='Cambria')

    if data.get('education'):
        _classic_section_heading(doc, lbl['education'])
        for edu in data['education']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
            _set_font(p.add_run(edu.get('school','')), 10, bold=True, font_name='Cambria')
            _set_font(p.add_run(f"\t{edu.get('graduation','')}"), 10, font_name='Cambria')
            deg = ', '.join(x for x in [edu.get('degree',''), edu.get('field','')] if x)
            if edu.get('gpa'): deg += f'  |  GPA: {edu["gpa"]}'
            if deg:
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(2)
                _set_font(p2.add_run(deg), 10, italic=True, font_name='Cambria')

    if data.get('skills'):
        _classic_section_heading(doc, lbl['skills'])
        for cat in data['skills']:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _set_font(p.add_run(f"{cat.get('category','')}: "), 10, bold=True, font_name='Cambria')
            _set_font(p.add_run(', '.join(cat.get('items', []))), 10, font_name='Cambria')

    langs = [l for l in data.get('languages', []) if l.get('language')]
    if langs:
        _classic_section_heading(doc, lbl['languages'])
        strs = [_lang_str(l) for l in langs]
        p = doc.add_paragraph()
        _set_font(p.add_run(', '.join(strs)), 10, font_name='Cambria')

    if data.get('projects'):
        _classic_section_heading(doc, lbl['projects'])
        for proj in data['projects']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            _set_font(p.add_run(f"{proj.get('name','')}"), 10, bold=True, font_name='Cambria')
            if proj.get('technologies'):
                _set_font(p.add_run(f" | {proj['technologies']}"), 10, italic=True, font_name='Cambria')
            if proj.get('description'):
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(2)
                _set_font(p2.add_run(proj['description']), 10, font_name='Cambria')

    certs = [c for c in data.get('certifications', []) if c]
    if certs:
        _classic_section_heading(doc, lbl['certifications'])
        for cert in certs:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            _set_font(p.add_run(f'— {cert}'), 10, font_name='Cambria')

    doc.save(output_path)


# ══════════════════════════════════════════════════════════════════════════════
#  MINIMAL TEMPLATE
# ══════════════════════════════════════════════════════════════════════════════

DRK2 = RGBColor(0x22, 0x22, 0x22)
GR2  = RGBColor(0x55, 0x55, 0x55)
GR3  = RGBColor(0x99, 0x99, 0x99)


def _minimal_section_heading(doc, title):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after  = Pt(4)
    _add_top_border(para, 'CCCCCC')
    run = para.add_run(title.upper())
    _set_font(run, 8, bold=True, color=GR3)
    return para


def _build_minimal_docx(data: dict, output_path: str, language: str = 'english'):
    doc = Document()
    _set_margins(doc, left=0.85, right=0.85)
    doc.styles['Normal'].paragraph_format.space_after  = Pt(0)
    lbl = _section_labels(language, 'minimal')

    _add_photo_header(doc, data, 20, DRK2, 9, GR3, center=False)

    if data.get('summary'):
        _minimal_section_heading(doc, lbl['summary'])
        p = doc.add_paragraph(data['summary'])
        _set_font(p.runs[0], 10, color=GR2)

    if data.get('experience'):
        _minimal_section_heading(doc, lbl['experience'])
        for exp in data['experience']:
            end = exp.get('end_date') or _present_label(language)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT)
            _set_font(p.add_run(exp.get('company','')), 10, bold=True, color=DRK2)
            _set_font(p.add_run(f"\t{exp.get('start_date','')} – {end}"), 9, color=GR3)
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(3)
            _set_font(p2.add_run(exp.get('title','')), 9, italic=True, color=GR2)
            for bullet in exp.get('bullets', []):
                bp = doc.add_paragraph()
                bp.paragraph_format.left_indent = Inches(0.15)
                bp.paragraph_format.space_after  = Pt(2)
                _set_font(bp.add_run(f'· {bullet}'), 9.5, color=GR2)

    if data.get('education'):
        _minimal_section_heading(doc, lbl['education'])
        for edu in data['education']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT)
            _set_font(p.add_run(edu.get('school','')), 10, bold=True, color=DRK2)
            _set_font(p.add_run(f"\t{edu.get('graduation','')}"), 9, color=GR3)
            deg = ', '.join(x for x in [edu.get('degree',''), edu.get('field','')] if x)
            if edu.get('gpa'): deg += f'  ·  GPA {edu["gpa"]}'
            if deg:
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(2)
                _set_font(p2.add_run(deg), 9, color=GR2)

    if data.get('skills'):
        _minimal_section_heading(doc, lbl['skills'])
        for cat in data['skills']:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _set_font(p.add_run(f"{cat.get('category','')}: "), 9.5, bold=True, color=DRK2)
            _set_font(p.add_run(', '.join(cat.get('items', []))), 9.5, color=GR2)

    langs = [l for l in data.get('languages', []) if l.get('language')]
    if langs:
        _minimal_section_heading(doc, lbl['languages'])
        strs = [_lang_str(l) for l in langs]
        p = doc.add_paragraph()
        _set_font(p.add_run(', '.join(strs)), 9.5, color=GR2)

    if data.get('projects'):
        _minimal_section_heading(doc, lbl['projects'])
        for proj in data['projects']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            _set_font(p.add_run(proj.get('name','')), 9.5, bold=True, color=DRK2)
            if proj.get('technologies'):
                _set_font(p.add_run(f" · {proj['technologies']}"), 9, color=GR3)
            if proj.get('description'):
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(2)
                _set_font(p2.add_run(proj['description']), 9.5, color=GR2)

    certs = [c for c in data.get('certifications', []) if c]
    if certs:
        _minimal_section_heading(doc, lbl['certifications'])
        for cert in certs:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _set_font(p.add_run(f'· {cert}'), 9.5, color=GR2)

    doc.save(output_path)


# ══════════════════════════════════════════════════════════════════════════════
#  JAPANESE TEMPLATE  (職務経歴書)
# ══════════════════════════════════════════════════════════════════════════════

JBL  = RGBColor(0x00, 0x33, 0x66)   # deep blue
JGR  = RGBColor(0x33, 0x33, 0x33)
JLT  = RGBColor(0x66, 0x66, 0x66)

JP_SECTIONS = {
    'summary':        '志望動機',
    'experience':     '職歴',
    'education':      '学歴',
    'skills':         '特技・スキル',
    'languages':      '語学力',
    'projects':       '主なプロジェクト',
    'certifications': '免許・資格',
    'extra':          '自己PR',
}


def _jp_section_heading(doc, title):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(3)
    run = para.add_run(f'■ {title}')
    _set_font(run, 11, bold=True, color=JBL, font_name='Yu Gothic')
    _add_bottom_border(para, '003366')
    return para


def _build_japanese_docx(data: dict, output_path: str, language: str = 'japanese'):
    doc = Document()
    _set_margins(doc, top=0.6, bottom=0.6, left=0.8, right=0.8)
    doc.styles['Normal'].paragraph_format.space_after  = Pt(0)

    # Document title + optional photo
    import datetime as _dt
    today    = _dt.date.today()
    date_str = f'作成日：{today.year}年{today.month}月{today.day}日'
    stream   = _photo_stream(data)

    if stream:
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OE
        table = doc.add_table(rows=1, cols=2)
        for cell in table._cells:
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            bdr = _OE('w:tcBorders')
            for edge in ('top','left','bottom','right','insideH','insideV'):
                tag = _OE(f'w:{edge}'); tag.set(_qn('w:val'), 'none'); bdr.append(tag)
            tcPr.append(bdr)
        table.columns[0].width = Inches(5.2)
        table.columns[1].width = Inches(1.6)

        left  = table.cell(0, 0)
        right = table.cell(0, 1)

        tp = left.paragraphs[0]
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(tp.add_run('履歴書'), 20, bold=True, color=JBL, font_name='Yu Gothic')

        rp = left.add_paragraph()
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rp.paragraph_format.space_after = Pt(4)
        rp.paragraph_format.tab_stops.add_tab_stop(Inches(4.8), WD_TAB_ALIGNMENT.RIGHT)
        _set_font(rp.add_run(data.get('name', '')), 13, bold=True, font_name='Yu Gothic')
        _set_font(rp.add_run(f'\t{date_str}'), 9, color=JLT, font_name='Yu Gothic')

        contact_parts = [data.get('email',''), data.get('phone',''), data.get('location','')]
        contact = '　｜　'.join(p for p in contact_parts if p)
        if contact:
            cp = left.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(6)
            _set_font(cp.add_run(contact), 9, color=JGR, font_name='Yu Gothic')

        # Photo cell
        pp = right.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_before = Pt(4)
        try:
            pp.add_run().add_picture(stream, width=Inches(1.3))
        except Exception:
            pass
    else:
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_after = Pt(2)
        _set_font(title_p.add_run('履歴書'), 20, bold=True, color=JBL, font_name='Yu Gothic')

        row_p = doc.add_paragraph()
        row_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_p.paragraph_format.space_after = Pt(8)
        row_p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        _set_font(row_p.add_run(data.get('name', '')), 13, bold=True, font_name='Yu Gothic')
        _set_font(row_p.add_run(f'\t{date_str}'), 9, color=JLT, font_name='Yu Gothic')

        contact_parts = [data.get('email',''), data.get('phone',''), data.get('location','')]
        contact = '　｜　'.join(p for p in contact_parts if p)
        if contact:
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(10)
            _set_font(cp.add_run(contact), 9, color=JGR, font_name='Yu Gothic')

    lbl = _section_labels(language, 'japanese')

    # ── 学歴（履歴書では最初）────────────────────────────────────────────────
    if data.get('education'):
        _jp_section_heading(doc, lbl['education'])
        for edu in data['education']:
            deg = '　'.join(x for x in [edu.get('degree',''), edu.get('field','')] if x)
            grad = edu.get('graduation','')
            school = edu.get('school','')
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(1)
            _set_font(p.add_run(f"{school}　{deg}　卒業"), 10, font_name='Yu Gothic')
            if grad:
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(2)
                _set_font(p2.add_run(f"　　卒業年月：{grad}"), 9, color=JGR, font_name='Yu Gothic')

    # ── 職歴 ────────────────────────────────────────────────────────────────
    if data.get('experience'):
        _jp_section_heading(doc, lbl['experience'])
        exps = data['experience']
        for i, exp in enumerate(exps):
            is_last = (i == len(exps) - 1)
            end_raw = exp.get('end_date', '') or ''
            is_current = not end_raw or end_raw.lower() in ('present', '現在', '至今', '')
            end_label = '現在に至る' if is_current else end_raw

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(7)
            _set_font(p.add_run(f"◆ {exp.get('company','')}　入社"), 11, bold=True, color=JBL, font_name='Yu Gothic')
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(2)
            _set_font(p2.add_run(f"　職位：{exp.get('title','')}　　{exp.get('start_date','')} 〜 {end_label}"), 9, color=JGR, font_name='Yu Gothic')
            for bullet in exp.get('bullets', []):
                bp = doc.add_paragraph()
                bp.paragraph_format.left_indent = Inches(0.2)
                bp.paragraph_format.space_after  = Pt(1)
                _set_font(bp.add_run(f'・{bullet}'), 10, font_name='Yu Gothic')
            if not is_current:
                pe = doc.add_paragraph()
                pe.paragraph_format.space_after = Pt(2)
                _set_font(pe.add_run(f"　同社　退職"), 9, color=JGR, font_name='Yu Gothic')
            elif is_last:
                pe = doc.add_paragraph()
                pe.paragraph_format.space_after = Pt(2)
                _set_font(pe.add_run('　　　　　　　　　　　　　　　　　　以上'), 9, color=JGR, font_name='Yu Gothic')

    # ── 免許・資格 ──────────────────────────────────────────────────────────
    certs = [c for c in data.get('certifications', []) if c]
    if certs:
        _jp_section_heading(doc, lbl['certifications'])
        for cert in certs:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            _set_font(p.add_run(f'・{cert}'), 10, font_name='Yu Gothic')

    # ── 志望動機 ────────────────────────────────────────────────────────────
    if data.get('summary'):
        _jp_section_heading(doc, lbl['summary'])
        p = doc.add_paragraph(data['summary'])
        _set_font(p.runs[0], 10, font_name='Yu Gothic')

    # ── 特技・スキル ────────────────────────────────────────────────────────
    if data.get('skills'):
        _jp_section_heading(doc, lbl['skills'])
        for cat in data['skills']:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _set_font(p.add_run(f"{cat.get('category','')}: "), 10, bold=True, font_name='Yu Gothic')
            _set_font(p.add_run(', '.join(cat.get('items', []))), 10, font_name='Yu Gothic')

    # ── 語学力 ──────────────────────────────────────────────────────────────
    langs = [l for l in data.get('languages', []) if l.get('language')]
    if langs:
        _jp_section_heading(doc, lbl['languages'])
        for l in langs:
            prof = l.get('proficiency', '')
            cert = l.get('certificate', '').strip()
            text = f"・{l['language']}：{prof}" if prof else f"・{l['language']}"
            if cert:
                text += f"（{cert}）"
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _set_font(p.add_run(text), 10, font_name='Yu Gothic')

    # ── 主なプロジェクト ────────────────────────────────────────────────────
    if data.get('projects'):
        _jp_section_heading(doc, lbl['projects'])
        for proj in data['projects']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            _set_font(p.add_run(proj.get('name','')), 10, bold=True, font_name='Yu Gothic')
            if proj.get('technologies'):
                _set_font(p.add_run(f"　({proj['technologies']})"), 9, color=JLT, font_name='Yu Gothic')
            if proj.get('description'):
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(2)
                _set_font(p2.add_run(proj['description']), 10, font_name='Yu Gothic')

    doc.save(output_path)


# ══════════════════════════════════════════════════════════════════════════════
#  TAIWANESE TEMPLATE (繁體中文履歷表)
# ══════════════════════════════════════════════════════════════════════════════

TWR  = RGBColor(0x7B, 0x14, 0x14)   # deep crimson accent
TWGR = RGBColor(0x33, 0x33, 0x33)
TWLT = RGBColor(0x66, 0x66, 0x66)

TW_SECTIONS = {
    'summary':        '個人簡介',
    'experience':     '工作經歷',
    'education':      '學歷',
    'skills':         '專業技能',
    'languages':      '語言能力',
    'projects':       '專案經歷',
    'certifications': '證照與認證',
}

TW_FONT = 'Microsoft JhengHei'   # Traditional Chinese; falls back gracefully on macOS


def _tw_section_heading(doc, title):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(3)
    run = para.add_run(f'■ {title}')
    _set_font(run, 11, bold=True, color=TWR, font_name=TW_FONT)
    _add_bottom_border(para, '7B1414')
    return para


def _build_taiwanese_docx(data: dict, output_path: str, language: str = 'taiwanese'):
    doc = Document()
    _set_margins(doc, top=0.6, bottom=0.6, left=0.8, right=0.8)
    doc.styles['Normal'].paragraph_format.space_after = Pt(0)

    import datetime as _dt
    today    = _dt.date.today()
    date_str = f'製作日期：{today.year}年{today.month}月{today.day}日'
    stream   = _photo_stream(data)

    if stream:
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OE
        table = doc.add_table(rows=1, cols=2)
        for cell in table._cells:
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            bdr = _OE('w:tcBorders')
            for edge in ('top','left','bottom','right','insideH','insideV'):
                tag = _OE(f'w:{edge}'); tag.set(_qn('w:val'), 'none'); bdr.append(tag)
            tcPr.append(bdr)
        table.columns[0].width = Inches(5.2)
        table.columns[1].width = Inches(1.6)

        left  = table.cell(0, 0)
        right = table.cell(0, 1)

        tp = left.paragraphs[0]
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(tp.add_run('履歷表'), 20, bold=True, color=TWR, font_name=TW_FONT)

        rp = left.add_paragraph()
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rp.paragraph_format.space_after = Pt(4)
        rp.paragraph_format.tab_stops.add_tab_stop(Inches(4.8), WD_TAB_ALIGNMENT.RIGHT)
        _set_font(rp.add_run(data.get('name', '')), 13, bold=True, font_name=TW_FONT)
        _set_font(rp.add_run(f'\t{date_str}'), 9, color=TWLT, font_name=TW_FONT)

        contact_parts = [data.get('email',''), data.get('phone',''), data.get('location','')]
        contact = '　｜　'.join(p for p in contact_parts if p)
        if contact:
            cp = left.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(6)
            _set_font(cp.add_run(contact), 9, color=TWGR, font_name=TW_FONT)

        pp = right.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_before = Pt(4)
        try:
            pp.add_run().add_picture(stream, width=Inches(1.3))
        except Exception:
            pass
    else:
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_after = Pt(2)
        _set_font(title_p.add_run('履歷表'), 20, bold=True, color=TWR, font_name=TW_FONT)

        row_p = doc.add_paragraph()
        row_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_p.paragraph_format.space_after = Pt(8)
        row_p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        _set_font(row_p.add_run(data.get('name', '')), 13, bold=True, font_name=TW_FONT)
        _set_font(row_p.add_run(f'\t{date_str}'), 9, color=TWLT, font_name=TW_FONT)

        contact_parts = [data.get('email',''), data.get('phone',''), data.get('location','')]
        contact = '　｜　'.join(p for p in contact_parts if p)
        if contact:
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(10)
            _set_font(cp.add_run(contact), 9, color=TWGR, font_name=TW_FONT)

    lbl = _section_labels(language, 'taiwanese')

    if data.get('summary'):
        _tw_section_heading(doc, lbl['summary'])
        p = doc.add_paragraph(data['summary'])
        _set_font(p.runs[0], 10, font_name=TW_FONT)

    if data.get('experience'):
        _tw_section_heading(doc, lbl['experience'])
        for exp in data['experience']:
            end = exp.get('end_date') or '至今'
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(7)
            _set_font(p.add_run(f"◆ {exp.get('company','')}"), 11, bold=True, color=TWR, font_name=TW_FONT)
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(2)
            _set_font(p2.add_run(f"　職稱：{exp.get('title','')}　　期間：{exp.get('start_date','')} 〜 {end}"), 9, color=TWGR, font_name=TW_FONT)
            for bullet in exp.get('bullets', []):
                bp = doc.add_paragraph()
                bp.paragraph_format.left_indent = Inches(0.2)
                bp.paragraph_format.space_after  = Pt(1)
                _set_font(bp.add_run(f'・{bullet}'), 10, font_name=TW_FONT)

    if data.get('skills'):
        _tw_section_heading(doc, lbl['skills'])
        for cat in data['skills']:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _set_font(p.add_run(f"{cat.get('category','')}: "), 10, bold=True, font_name=TW_FONT)
            _set_font(p.add_run(', '.join(cat.get('items', []))), 10, font_name=TW_FONT)

    langs = [l for l in data.get('languages', []) if l.get('language')]
    if langs:
        _tw_section_heading(doc, lbl['languages'])
        for l in langs:
            prof = l.get('proficiency', '')
            cert = l.get('certificate', '').strip()
            text = f"・{l['language']}：{prof}" if prof else f"・{l['language']}"
            if cert:
                text += f"（{cert}）"
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _set_font(p.add_run(text), 10, font_name=TW_FONT)

    if data.get('education'):
        _tw_section_heading(doc, lbl['education'])
        for edu in data['education']:
            deg = '、'.join(x for x in [edu.get('degree',''), edu.get('field','')] if x)
            grad = edu.get('graduation','')
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(2)
            _set_font(p.add_run(f"{edu.get('school','')}　{deg}　{grad}畢業"), 10, font_name=TW_FONT)

    if data.get('projects'):
        _tw_section_heading(doc, lbl['projects'])
        for proj in data['projects']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            _set_font(p.add_run(proj.get('name','')), 10, bold=True, font_name=TW_FONT)
            if proj.get('technologies'):
                _set_font(p.add_run(f"　({proj['technologies']})"), 9, color=TWLT, font_name=TW_FONT)
            if proj.get('description'):
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(2)
                _set_font(p2.add_run(proj['description']), 10, font_name=TW_FONT)

    certs = [c for c in data.get('certifications', []) if c]
    if certs:
        _tw_section_heading(doc, lbl['certifications'])
        for cert in certs:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            _set_font(p.add_run(f'・{cert}'), 10, font_name=TW_FONT)

    doc.save(output_path)


# ══════════════════════════════════════════════════════════════════════════════
#  PDF — template-aware HTML renderer
# ══════════════════════════════════════════════════════════════════════════════

_TEMPLATE_CSS = {
    'modern': """
      body { font-family: -apple-system, 'Segoe UI', sans-serif; font-size: 10pt; color: #222; }
      .name { font-size: 22pt; font-weight: 700; text-align: center; color: #1a1a2e; margin-bottom: 4px; }
      .contact { text-align: center; color: #555; font-size: 9pt; margin-bottom: 14px; }
      .section-title { font-size: 9.5pt; font-weight: 700; text-transform: uppercase; letter-spacing: .06em;
        color: #1a1a2e; border-bottom: 1.5px solid #1a1a2e; padding-bottom: 2px; margin: 12px 0 6px; }
      .entry-header { display: flex; justify-content: space-between; }
      .bold { font-weight: 700; } .date { font-size: 9pt; color: #666; }
      .subtitle { font-style: italic; color: #444; margin-bottom: 3px; }
      ul { padding-left: 16px; margin: 3px 0; } li { margin-bottom: 2px; }
      p { margin-bottom: 3px; } .tech { color: #224488; }
    """,
    'classic': """
      body { font-family: Georgia, 'Times New Roman', serif; font-size: 10.5pt; color: #111; }
      .name { font-size: 18pt; font-weight: 700; text-align: center; color: #000; margin-bottom: 4px; }
      .contact { text-align: center; color: #444; font-size: 9.5pt; margin-bottom: 16px; }
      .section-title { font-size: 11pt; font-weight: 700; text-transform: uppercase; text-decoration: underline;
        color: #000; margin: 14px 0 5px; }
      .entry-header { display: flex; justify-content: space-between; }
      .bold { font-weight: 700; } .date { font-size: 9.5pt; color: #333; }
      .subtitle { font-style: italic; color: #333; margin-bottom: 4px; }
      ul { list-style: none; padding-left: 16px; margin: 3px 0; }
      li::before { content: "— "; } li { margin-bottom: 3px; }
      p { margin-bottom: 4px; } .tech { font-style: italic; }
    """,
    'minimal': """
      body { font-family: 'Helvetica Neue', Arial, sans-serif; font-size: 10pt; color: #555; }
      .name { font-size: 20pt; font-weight: 300; color: #222; margin-bottom: 3px; }
      .contact { color: #999; font-size: 8.5pt; margin-bottom: 16px; }
      .section-title { font-size: 8pt; font-weight: 700; text-transform: uppercase; letter-spacing: .08em;
        color: #aaa; border-top: 1px solid #ddd; padding-top: 6px; margin: 14px 0 6px; }
      .entry-header { display: flex; justify-content: space-between; }
      .bold { font-weight: 600; color: #333; } .date { font-size: 8.5pt; color: #bbb; }
      .subtitle { font-style: italic; color: #777; margin-bottom: 3px; }
      ul { list-style: none; padding-left: 14px; margin: 3px 0; }
      li::before { content: "· "; color: #bbb; } li { margin-bottom: 3px; color: #555; }
      p { margin-bottom: 3px; } .tech { color: #aaa; }
    """,
    'japanese': """
      body { font-family: 'Hiragino Kaku Gothic Pro', 'Yu Gothic', 'Meiryo', sans-serif;
             font-size: 10pt; color: #222; }
      .jp-title { font-size: 20pt; font-weight: 700; text-align: center; color: #003366; margin-bottom: 4px; }
      .jp-name-row { display: flex; justify-content: space-between; align-items: baseline; margin-bottom: 4px; }
      .jp-name { font-size: 13pt; font-weight: 700; }
      .jp-date { font-size: 9pt; color: #666; }
      .contact { text-align: center; color: #555; font-size: 9pt; margin-bottom: 14px; }
      .section-title { font-size: 11pt; font-weight: 700; color: #003366;
        border-bottom: 2px solid #003366; padding-bottom: 2px; margin: 12px 0 6px; }
      .company { font-size: 11pt; font-weight: 700; color: #003366; margin-top: 8px; }
      .jp-role { font-size: 9pt; color: #555; margin-bottom: 3px; }
      .bold { font-weight: 700; } .date { font-size: 9pt; color: #666; }
      .subtitle { font-style: italic; color: #444; margin-bottom: 3px; }
      ul { list-style: none; padding-left: 12px; margin: 3px 0; }
      li::before { content: "・"; } li { margin-bottom: 2px; }
      p { margin-bottom: 3px; } .tech { color: #336; }
      .entry-header { display: flex; justify-content: space-between; }
    """,
    'taiwanese': """
      body { font-family: 'PingFang TC', 'Microsoft JhengHei', 'Heiti TC', 'Noto Sans TC', sans-serif;
             font-size: 10pt; color: #222; }
      .tw-title { font-size: 20pt; font-weight: 700; text-align: center; color: #7B1414; margin-bottom: 4px; }
      .tw-name-row { display: flex; justify-content: space-between; align-items: baseline; margin-bottom: 4px; }
      .tw-name { font-size: 13pt; font-weight: 700; }
      .tw-date { font-size: 9pt; color: #666; }
      .contact { text-align: center; color: #555; font-size: 9pt; margin-bottom: 14px; }
      .section-title { font-size: 11pt; font-weight: 700; color: #7B1414;
        border-bottom: 2px solid #7B1414; padding-bottom: 2px; margin: 12px 0 6px; }
      .company { font-size: 11pt; font-weight: 700; color: #7B1414; margin-top: 8px; }
      .tw-role { font-size: 9pt; color: #555; margin-bottom: 3px; }
      .bold { font-weight: 700; } .date { font-size: 9pt; color: #666; }
      .subtitle { font-style: italic; color: #444; margin-bottom: 3px; }
      ul { list-style: none; padding-left: 12px; margin: 3px 0; }
      li::before { content: "・"; } li { margin-bottom: 2px; }
      p { margin-bottom: 3px; } .tech { color: #7B1414; }
      .entry-header { display: flex; justify-content: space-between; }
    """,
}


def _render_html(data: dict, template: str = 'modern', language: str = 'english') -> str:
    is_japanese_tmpl  = template == 'japanese'
    is_taiwanese_tmpl = template == 'taiwanese'
    is_japanese  = is_japanese_tmpl
    is_taiwanese = is_taiwanese_tmpl
    is_cjk_lang  = language in ('japanese', 'taiwanese')
    is_cjk       = is_japanese or is_taiwanese or is_cjk_lang
    css = _TEMPLATE_CSS.get(template, _TEMPLATE_CSS['modern'])

    contact_parts = [data.get('email',''), data.get('phone',''), data.get('location',''),
                     data.get('linkedin',''), data.get('github','')]
    sep = '　｜　' if is_cjk else ' &nbsp;|&nbsp; '
    contact_line = sep.join(p for p in contact_parts if p)

    sections = []

    photo     = data.get('photo', '')
    photo_tag = f'<img src="{photo}" style="width:1in;height:1.2in;object-fit:cover;object-position:center top;border:1px solid #ccc;border-radius:3px;flex-shrink:0">' if photo else ''

    import datetime
    today = datetime.date.today()
    _photo_el = (
        f'<div style="flex-shrink:0;text-align:center">'
        f'<img src="{photo}" style="width:1in;height:1.25in;object-fit:cover;object-position:center top;'
        f'border:1px solid #ccc;border-radius:3px;display:block"></div>'
    ) if photo else ''

    if is_japanese:
        date_str = f'作成日：{today.year}年{today.month}月{today.day}日'
        if photo:
            header = f'''
              <div style="display:flex;align-items:flex-start;gap:20px;margin-bottom:6px">
                <div style="flex:1;min-width:0;text-align:center">
                  <div class="jp-title">履歴書</div>
                  <div style="font-size:13pt;font-weight:700;margin:4px 0">{data.get("name","")}</div>
                  <div style="font-size:9pt;color:#666;margin-bottom:5px">{date_str}</div>
                  <div class="contact">{contact_line}</div>
                </div>
                {_photo_el}
              </div>'''
        else:
            header = f'''
              <div class="jp-title">履歴書</div>
              <div class="jp-name-row">
                <span class="jp-name">{data.get("name","")}</span>
                <span class="jp-date">{date_str}</span>
              </div>
              <div class="contact" style="margin-bottom:10px">{contact_line}</div>'''
        labels = _section_labels(language, 'japanese')
    elif is_taiwanese:
        date_str = f'製作日期：{today.year}年{today.month}月{today.day}日'
        if photo:
            header = f'''
              <div style="display:flex;align-items:flex-start;gap:20px;margin-bottom:6px">
                <div style="flex:1;min-width:0;text-align:center">
                  <div class="tw-title">履歷表</div>
                  <div style="font-size:13pt;font-weight:700;margin:4px 0">{data.get("name","")}</div>
                  <div style="font-size:9pt;color:#666;margin-bottom:5px">{date_str}</div>
                  <div class="contact">{contact_line}</div>
                </div>
                {_photo_el}
              </div>'''
        else:
            header = f'''
              <div class="tw-title">履歷表</div>
              <div class="tw-name-row">
                <span class="tw-name">{data.get("name","")}</span>
                <span class="tw-date">{date_str}</span>
              </div>
              <div class="contact" style="margin-bottom:10px">{contact_line}</div>'''
        labels = _section_labels(language, 'taiwanese')
    else:
        name_contact = f'''
          <div class="name">{data.get("name","")}</div>
          <div class="contact">{contact_line}</div>'''
        if photo:
            photo_el = f'<img src="{photo}" style="width:0.9in;height:1.1in;object-fit:cover;object-position:center top;border-radius:4px;flex-shrink:0">'
            header = f'<div style="display:flex;align-items:flex-start;justify-content:space-between;gap:16px"><div style="flex:1">{name_contact}</div>{photo_el}</div>'
        else:
            header = name_contact
        labels = _section_labels(language, template)

    prefix = '■ ' if is_cjk else ''

    # ── build reusable section fragments ─────────────────────────────────────

    def _summary_sec():
        if not data.get('summary'): return
        sections.append(f'<div class="section-title">{prefix}{labels["summary"]}</div><p>{data["summary"]}</p>')

    def _experience_sec():
        if not data.get('experience'): return
        exp_html = ''
        exps = data['experience']
        for i, exp in enumerate(exps):
            is_last = (i == len(exps) - 1)
            end_raw = exp.get('end_date', '') or ''
            is_current = not end_raw or end_raw.lower() in ('present', '現在', '至今', '')
            end = _present_label(language) if is_current else end_raw
            bullets = ''.join(f'<li>{b}</li>' for b in exp.get('bullets', []))
            if is_japanese:
                exp_html += f'''
                <div class="company">◆ {exp.get("company","")}　入社</div>
                <div class="jp-role">職位：{exp.get("title","")}　{exp.get("start_date","")} 〜 {end}</div>
                <ul>{bullets}</ul>'''
                if not is_current:
                    exp_html += f'<div class="jp-role" style="text-align:right">同社　退職</div>'
                elif is_last:
                    exp_html += f'<div class="jp-role" style="text-align:right">現在に至る</div>'
            elif is_taiwanese:
                exp_html += f'''
                <div class="company">◆ {exp.get("company","")}</div>
                <div class="tw-role">職稱：{exp.get("title","")}　期間：{exp.get("start_date","")} 〜 {end}</div>
                <ul>{bullets}</ul>'''
            else:
                exp_html += f'''
                <div class="entry">
                  <div class="entry-header">
                    <span class="bold">{exp.get("company","")}</span>
                    <span class="date">{exp.get("start_date","")} – {end}</span>
                  </div>
                  <div class="subtitle">{exp.get("title","")}</div>
                  <ul>{bullets}</ul>
                </div>'''
        sections.append(f'<div class="section-title">{prefix}{labels["experience"]}</div>{exp_html}')

    def _skills_sec():
        if not data.get('skills'): return
        skill_html = ''.join(
            f'<p><strong>{c.get("category","")}: </strong>{", ".join(c.get("items",[]))}</p>'
            for c in data['skills']
        )
        sections.append(f'<div class="section-title">{prefix}{labels["skills"]}</div>{skill_html}')

    def _languages_sec():
        langs = [l for l in data.get('languages', []) if l.get('language')]
        if not langs: return
        if is_cjk:
            def _cjk_lang(l):
                prof = l.get('proficiency', '')
                cert = l.get('certificate', '').strip()
                base = f"・{l['language']}：{prof}" if prof else f"・{l['language']}"
                return f'<p>{base}{"（" + cert + "）" if cert else ""}</p>'
            lang_html = ''.join(_cjk_lang(l) for l in langs)
        else:
            lang_html = f'<p>{", ".join(_lang_str(l) for l in langs)}</p>'
        sections.append(f'<div class="section-title">{prefix}{labels["languages"]}</div>{lang_html}')

    def _education_sec():
        if not data.get('education'): return
        edu_html = ''
        for edu in data['education']:
            deg = ', '.join(x for x in [edu.get('degree',''), edu.get('field','')] if x)
            if edu.get('gpa'): deg += f' | GPA: {edu["gpa"]}'
            if is_japanese:
                edu_html += f'<p><strong>{edu.get("school","")}</strong>　{deg}　{edu.get("graduation","")}卒業</p>'
            elif is_taiwanese:
                edu_html += f'<p><strong>{edu.get("school","")}</strong>　{deg}　{edu.get("graduation","")}畢業</p>'
            else:
                edu_html += f'''
                <div class="entry">
                  <div class="entry-header">
                    <span class="bold">{edu.get("school","")}</span>
                    <span class="date">{edu.get("graduation","")}</span>
                  </div>
                  <div class="subtitle">{deg}</div>
                </div>'''
        sections.append(f'<div class="section-title">{prefix}{labels["education"]}</div>{edu_html}')

    def _projects_sec():
        if not data.get('projects'): return
        proj_html = ''
        for proj in data['projects']:
            tech = f' <span class="tech">({proj["technologies"]})</span>' if proj.get('technologies') else ''
            proj_html += f'<div class="entry"><p><strong>{proj.get("name","")}</strong>{tech}</p>'
            if proj.get('description'): proj_html += f'<p>{proj["description"]}</p>'
            proj_html += '</div>'
        sections.append(f'<div class="section-title">{prefix}{labels["projects"]}</div>{proj_html}')

    def _certs_sec():
        certs = [c for c in data.get('certifications', []) if c]
        if not certs: return
        cert_html = ''.join(f'<li>{c}</li>' for c in certs)
        sections.append(f'<div class="section-title">{prefix}{labels["certifications"]}</div><ul>{cert_html}</ul>')

    # ── section order ─────────────────────────────────────────────────────────
    if is_japanese:
        # 履歴書 order: 学歴 → 職歴 → 免許・資格 → 志望動機 → スキル → 語学力 → プロジェクト
        _education_sec(); _experience_sec(); _certs_sec()
        _summary_sec(); _skills_sec(); _languages_sec(); _projects_sec()
    else:
        # Western & Taiwanese order: summary → experience → skills → languages → education → projects → certs
        _summary_sec(); _experience_sec(); _skills_sec()
        _languages_sec(); _education_sec(); _projects_sec(); _certs_sec()

    return f'''<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{ padding: 0.5in 0.75in; line-height: 1.4; }}
  .entry {{ margin-bottom: 8px; }}
  {css}
</style>
</head>
<body>
  {header}
  {"".join(f'<div class="section">{s}</div>' for s in sections)}
</body>
</html>'''


# ══════════════════════════════════════════════════════════════════════════════
#  COVER LETTER BUILDERS
# ══════════════════════════════════════════════════════════════════════════════

def build_cover_letter_docx(letter_data: dict, candidate: dict, output_path: str):
    doc = Document()
    _set_margins(doc, top=1.0, bottom=1.0, left=1.2, right=1.2)
    doc.styles['Normal'].paragraph_format.space_after = Pt(0)

    # Sender name
    p = doc.add_paragraph()
    _set_font(p.add_run(candidate.get('name', '')), 14, bold=True, color=DARK)
    contact = _contact_str(candidate)
    if contact:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(18)
        _set_font(p2.add_run(contact), 9, color=MID)

    # Body paragraphs
    for para_text in letter_data.get('text', '').split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            _set_font(p.add_run(para_text), 11)

    # Sign-off
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    _set_font(p.add_run('Sincerely,'), 11)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(6)
    _set_font(p2.add_run(candidate.get('name', '')), 11, bold=True)

    doc.save(output_path)


def build_cover_letter_pdf(letter_data: dict, candidate: dict, output_path: str):
    contact_parts = [candidate.get('email',''), candidate.get('phone',''),
                     candidate.get('location',''), candidate.get('linkedin','')]
    contact_line = ' | '.join(p for p in contact_parts if p)
    paragraphs = ''.join(
        f'<p>{para.strip()}</p>'
        for para in letter_data.get('text', '').split('\n\n')
        if para.strip()
    )
    html = f'''<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
  body {{ font-family: Georgia, serif; font-size: 11pt; color: #222;
         padding: 1in 1.2in; line-height: 1.7; max-width: 700px; }}
  .sender-name {{ font-size: 15pt; font-weight: bold; color: #1a1a2e; margin-bottom: 4px; }}
  .contact {{ font-size: 9pt; color: #666; margin-bottom: 28px; }}
  p {{ margin-bottom: 14px; }}
  .signoff {{ margin-top: 20px; }}
  .sig-name {{ font-weight: bold; margin-top: 8px; }}
</style></head><body>
  <div class="sender-name">{candidate.get("name","")}</div>
  <div class="contact">{contact_line}</div>
  {paragraphs}
  <div class="signoff">Sincerely,<div class="sig-name">{candidate.get("name","")}</div></div>
</body></html>'''
    weasyprint.HTML(string=html).write_pdf(output_path)


def build_cover_letter_md(letter_data: dict, candidate: dict, output_path: str):
    lines = [
        f"# Cover Letter — {candidate.get('name', '')}",
        '',
        _contact_str(candidate),
        '',
        '---',
        '',
    ]
    for para in letter_data.get('text', '').split('\n\n'):
        if para.strip():
            lines.append(para.strip())
            lines.append('')
    lines += ['---', '', f'Sincerely,  ', f'**{candidate.get("name","")}**']
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))


# ══════════════════════════════════════════════════════════════════════════════
#  PUBLIC DISPATCHERS
# ══════════════════════════════════════════════════════════════════════════════

_DOCX_BUILDERS = {
    'modern':     _build_modern_docx,
    'classic':    _build_classic_docx,
    'minimal':    _build_minimal_docx,
    'japanese':   _build_japanese_docx,
    'taiwanese':  _build_taiwanese_docx,
}


def build_docx(data: dict, output_path: str, template: str = 'modern', language: str = 'english'):
    _DOCX_BUILDERS.get(template, _build_modern_docx)(data, output_path, language)


def build_pdf(data: dict, output_path: str, template: str = 'modern', language: str = 'english'):
    weasyprint.HTML(string=_render_html(data, template, language)).write_pdf(output_path)


def build_preview_html(data: dict, output_path: str, template: str = 'modern', language: str = 'english'):
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(_render_html(data, template, language))


def build_markdown(data: dict, output_path: str):
    import datetime
    lines = [f"# {data.get('name', '')}\n"]
    contact = _contact_str(data)
    if contact: lines += [contact, '']

    def sec(title, content_lines):
        lines.append(f'\n## {title}\n')
        lines.extend(content_lines)
        lines.append('')

    if data.get('summary'): sec('Professional Summary', [data['summary']])

    if data.get('experience'):
        exp_lines = []
        for exp in data['experience']:
            end = exp.get('end_date') or 'Present'
            exp_lines.append(f"### {exp.get('company','')}  \n**{exp.get('title','')}** | {exp.get('start_date','')} – {end}")
            for b in exp.get('bullets', []): exp_lines.append(f"- {b}")
            exp_lines.append('')
        sec('Experience', exp_lines)

    if data.get('education'):
        edu_lines = []
        for edu in data['education']:
            deg = ', '.join(x for x in [edu.get('degree',''), edu.get('field','')] if x)
            gpa = f" | GPA: {edu['gpa']}" if edu.get('gpa') else ''
            edu_lines += [f"### {edu.get('school','')}", f"{deg}{gpa} | {edu.get('graduation','')}", '']
        sec('Education', edu_lines)

    if data.get('skills'):
        skill_lines = [f"**{c.get('category','')}:** {', '.join(c.get('items',[]))}" for c in data['skills']]
        sec('Skills', skill_lines)

    langs = [l for l in data.get('languages', []) if l.get('language')]
    if langs:
        strs = [_lang_str(l) for l in langs]
        sec('Languages', [', '.join(strs)])

    if data.get('projects'):
        proj_lines = []
        for p in data['projects']:
            tech = f" | {p['technologies']}" if p.get('technologies') else ''
            proj_lines += [f"### {p.get('name','')}{tech}", p.get('description',''), '']
        sec('Projects', proj_lines)

    certs = [c for c in data.get('certifications', []) if c]
    if certs: sec('Certifications / Professional Licenses', [f"- {c}" for c in certs])

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
